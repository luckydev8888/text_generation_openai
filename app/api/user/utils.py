import csv
from bs4 import BeautifulSoup
from docx import Document
import re
import PyPDF2
import json
from datetime import datetime
from app.mongo import get_db
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from difflib import SequenceMatcher
from flask import session
from unidecode import unidecode


def get_pdf_text(pdf_path):
    text = ""
    with open(pdf_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text


def clean_text(str):
    cite_pattern = r"【\d+†source】|【\d+:\d+†source】|【\d+:\d+†fuente】|【.*?】"
    return re.sub(cite_pattern, "", str)


def find_setencia_list(list):
    db = get_db()
    collection = db["sentencias"]

    json_data = []
    for item in list:
        doc = collection.find_one({"providencia": item})
        if doc:
            json_data.append(
                {
                    "providencia": doc["providencia"],
                    #'tipo': doc['tipo'],
                    #'ano': doc['ano'],
                    "fecha_sentencia": doc["fecha_sentencia"],
                    "derechos": doc["derechos"],  # se cambio 'derechos' por 'derechos'
                    "magistrado": doc["magistrado"],
                    #'fecha_publicada': doc['fecha_publicada'],
                    "expediente": doc["expediente"],
                    "url": doc["url"],
                }
            )
    return json_data


def html_to_text(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    return soup.get_text()


def add_html_to_docx(soup, doc):
    for element in soup:
        if element == "\n":
            continue
        if element.name == "p":
            paragraph = doc.add_paragraph(element.get_text())
            run = paragraph.runs[0]
            run.font.size = Pt(11.5)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "h1":
            paragraph = doc.add_heading(level=1)
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(20)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "h2":
            paragraph = doc.add_heading(level=2)
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(16)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "h3":
            paragraph = doc.add_heading(level=3)
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(14)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "h4":
            paragraph = doc.add_heading(level=3)
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(13)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "strong":
            run = doc.add_paragraph().add_run(element.get_text())
            paragraph_format = run.paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "em":
            run = doc.add_paragraph().add_run(element.get_text())
            run.italic = True
            run.font.size = Pt(11.5)
            paragraph_format = run.paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after
        elif element.name == "ul":
            list_items = element.find_all("li")
            for i, li in enumerate(list_items):
                paragraph = doc.add_paragraph(li.get_text(), style="ListBullet")
                run = paragraph.runs[0]
                run.font.size = Pt(11.5)
                paragraph_format = paragraph.paragraph_format
                if i == 0:
                    paragraph_format.space_before = Pt(6)
                if i == len(list_items) - 1:
                    paragraph_format.space_after = Pt(6)
        elif element.name == "ol":
            list_items = element.find_all("li")
            for i, li in enumerate(list_items):
                paragraph = doc.add_paragraph(li.get_text(), style="ListNumber")
                run = paragraph.runs[0]
                run.font.size = Pt(11.5)
                paragraph_format = paragraph.paragraph_format
                if i == 0:
                    paragraph_format.space_before = Pt(6)
                if i == len(list_items) - 1:
                    paragraph_format.space_after = Pt(6)

        elif isinstance(element, str):
            paragraph = doc.add_paragraph(element)
            run = paragraph.runs[0]
            run.font.size = Pt(11.5)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)  # 0.5 line space after


def create_docx_from_html(html_content, filename):
    soup = BeautifulSoup(html_content, "lxml")

    doc = Document()
    add_html_to_docx(soup.body.contents, doc)

    doc.save(filename)


def get_constitution(str):
    pattern = r"Artículo (?P<number>\d+)"  # Named group to capture the number
    matches = re.findall(pattern, str)
    print("Found:", matches)

    return matches


def proccess_code(codigo):
    # Eliminar acentos y convertir a minúsculas
    codigo = unidecode(codigo.lower())

    codigo = re.sub(
        r"[^A-Za-z0-9]", "-", codigo
    )  # Reemplazar caracteres no alfanuméricos por guiones
    codigo = re.sub(r"de-", "-", codigo)  # Eliminar 'de-' seguido de un guion
    codigo = re.sub(
        r"-+", "-", codigo
    )  # Reemplazar múltiples guiones consecutivos por un solo guion
    codigo = re.sub(
        r"-(\d{2})(\d{2})$", r"-\2", codigo
    )  # Convertir años de 4 dígitos a 2 dígitos si es necesario
    return codigo


def parse_date(date_str):
    return datetime.strptime(date_str, "%Y-%m-%d")


# Función para calcular la similitud entre dos textos
def es_similar(texto1, texto2):
    return SequenceMatcher(None, texto1, texto2).ratio()


# Lista de verbos comunes para identificar peticiones (en infinitivo o imperativo)
VERBOS_COMUNES = r"\b(solicitar|realizar|suministrar|ordenar|iniciar|conceder|proteger|amparar|otorgar|dictaminar|evaluar|exigir|pedir|valorar|determinar|tutelar)\b"


def buscar_patrones_en_texto(TutelaSentTemp):
    print(
        f"TutelaSentTemp contiene {len(TutelaSentTemp)} elementos para análisis de sentencias adjuntas"
    )

    # Verificar que TutelaSentTemp sea una lista válida
    if not isinstance(TutelaSentTemp, list) or not TutelaSentTemp:
        print("Error: TutelaSentTemp no es una lista válida o está vacía.")
        return []

    # Inicializar la lista de sentencias coincidentes
    salida1 = []

    # Conectar a la base de datos MongoDB
    db = get_db()
    collection = db["sentencias"]

    # Crear un pipeline que busque coincidencias de las sentencias en la base de datos
    pipeline = [
        {"$project": {"_id": 1, "providencia": 1, "fecha_sentencia": 1}},
        {"$sort": {"fecha_sentencia": -1}},
    ]
    # Iteramos directamente sobre el cursor sin convertirlo en lista
    for doc in collection.aggregate(pipeline):
        providencia_db = doc["providencia"].lower().strip()
        providencia_db_normalizada = re.sub(
            r"(\w+)-(\d+)/(\d+)", r"\1-\2-\3", providencia_db
        )

        for sentencia_adjunta in TutelaSentTemp:
            sentencia_adjunta_normalizada = (
                re.sub(r"(\w+)-(\d+)/(\d+)", r"\1-\2-\3", sentencia_adjunta)
                .lower()
                .strip()
            )
            if sentencia_adjunta_normalizada == providencia_db_normalizada:
                print(
                    f"Coincidencia encontrada: '{providencia_db_normalizada}' en documento ID: {doc['_id']}"
                )
                salida1.append(providencia_db_normalizada)

    print(f"Total de coincidencias encontradas salida1: {len(salida1)}")
    return salida1


# Función para calcular la similitud entre dos textos (debes definir esta función)
def es_similar(derecho_a, derecho_b):
    # Lógica de comparación, por ejemplo con Levenshtein o alguna métrica de similitud
    pass


def buscar_derechos_en_mongo(TutelaDerecTemp):
    print(
        f"TutelaDerecTemp contiene {len(TutelaDerecTemp)} elementos como entrada análisis derechos"
    )

    # Accedemos a la base de datos
    db = get_db()
    collection = db["sentencias"]

    # Preparamos el pipeline de búsqueda
    pipeline = [
        {"$project": {"_id": 1, "providencia": 1, "fecha_sentencia": 1, "derechos": 1}},
        {"$sort": {"fecha_sentencia": -1}},  # Ordenamos por la fecha de sentencia
    ]

    salida2 = []
    coincidencias_derechos = 0

    # Iteramos sobre los documentos en el resultado del pipeline
    for doc in collection.aggregate(pipeline):
        derechos_db = (
            doc.get("derechos", "").lower().strip()
        )  # Convertimos derechos a minúsculas
        derechos_db_normalizado = re.sub(
            r"(\w+)-(\d+)/(\d+)", r"\1-\2-\3", derechos_db
        )  # Normalizamos el formato

        for derecho in TutelaDerecTemp:
            derecho_normalizado = re.sub(
                r"(\w+)-(\d+)/(\d+)", r"\1-\2-\3", derecho.lower().strip()
            )

            # Verificamos la similitud entre el derecho en la base de datos y el derecho de la lista de entrada
            similitud = es_similar(derecho_normalizado, derechos_db_normalizado)

            if similitud > 0.6:
                coincidencias_derechos += 1
                print(
                    f"Similitud alta entre '{derecho_normalizado}' y '{derechos_db_normalizado}' (Similitud: {similitud})"
                )

                # Si se encuentra una coincidencia, agregamos el documento a la salida
                salida2.append(
                    {
                        "_id": doc["_id"],
                        "providencia": doc["providencia"],
                        "fecha_sentencia": doc["fecha_sentencia"],
                        "similitud": similitud,
                    }
                )

    # Ordenamos por similitud y limitamos a 100 resultados
    salida2 = sorted(salida2, key=lambda x: x["similitud"], reverse=True)[:100]

    return salida2


# Función para calcular la similitud entre dos textos (debes definir esta función)
def es_similar(texto_a, texto_b):
    # Lógica de comparación, por ejemplo con Levenshtein o alguna métrica de similitud
    pass


def buscar_peticiones_con_ids(salida2, TutelaPetTemp):
    print(
        f"TutelaPetTemp contiene {len(TutelaPetTemp)} elementos como entrada análisis ids"
    )

    if not TutelaPetTemp:
        print(
            "Error: No se encontraron los datos de derechos fundamentales invocados en la sesión."
        )
        return salida2  # Si no hay peticiones, retornamos `salida2` sin modificaciones

    # Conectar a la base de datos MongoDB
    db = get_db()
    collection = db["sentencias"]

    salida3 = []
    SentPetTempFiles = []

    # Recorrer los IDs obtenidos en salida2
    for doc_info in salida2:
        doc_id = doc_info["_id"]  # Asumimos que en salida2 cada documento tiene '_id'
        doc = collection.find_one({"_id": doc_id})
        if doc:
            texto = doc.get(
                "texto", ""
            ).lower()  # Texto completo del documento en minúsculas
            coincidencias_peticiones = 0

            print(f"Analizando documento con ID {doc_id}")

            # Comparar las peticiones con el texto del documento
            for peticion in TutelaPetTemp:
                peticion_normalizada = peticion.lower().strip()

                # Coincidencia directa
                if peticion_normalizada in texto:
                    coincidencias_peticiones += 1
                    print(
                        f"Coincidencia directa para petición: '{peticion_normalizada}' en documento ID: {doc_id}"
                    )

                # Coincidencia parcial utilizando similitud
                oraciones_documento = re.split(
                    r"[.!?]", texto
                )  # Dividir el documento en oraciones
                for oracion in oraciones_documento:
                    similitud = es_similar(peticion_normalizada, oracion.strip())
                    if similitud > 0.65:  # Umbral de similitud flexible
                        coincidencias_peticiones += 1
                        print(
                            f"Similitud entre '{peticion}' y '{oracion[:50]}...' en documento ID: {doc_id}, Similitud: {similitud}"
                        )

            if coincidencias_peticiones > 0:
                SentPetTempFiles.append((doc, coincidencias_peticiones))

    if SentPetTempFiles:
        # Ordenamos los resultados por el número de coincidencias y limitamos el número de resultados
        SentPetTempFiles = sorted(SentPetTempFiles, key=lambda x: x[1], reverse=True)[
            :5
        ]
        salida3 = salida2 + [doc["providencia"] for doc, _ in SentPetTempFiles]

    return salida3


def get_sentencia(TutelaDerecTemp, TutelaPetTemp, TutelaSentTemp):
    print("Iniciando get_sentencia...")
    print(f"TutelaDerecTemp: {TutelaDerecTemp}")
    print(f"TutelaPetTemp: {TutelaPetTemp}")
    print(f"TutelaSentTemp: {TutelaSentTemp}")

    # Parte 1: Búsqueda de patrones en el texto (pattern matching)
    salida1 = buscar_patrones_en_texto(TutelaSentTemp)
    print(f"Resultados de buscar_patrones_en_texto: {salida1}")

    # Parte 2: Búsqueda de derechos en la base de datos (TutelaDerecTemp)
    salida2 = buscar_derechos_en_mongo(TutelaDerecTemp)
    print(f"Resultados de buscar_derechos_en_mongo: Salida2: {salida2}")

    # Parte 3: Búsqueda en Mongo de los IDs obtenidos y matching con TutelaPetTemp
    salida3 = buscar_peticiones_con_ids(salida2, salida2, TutelaPetTemp)
    print(f"Resultados de buscar_peticiones_con_ids: {salida3}")

    # Combinar las salidas sin duplicar elementos
    sentencia_list = []
    for salida in [salida1, salida2, salida3]:
        for item in salida:
            if item not in sentencia_list:
                sentencia_list.append(item)

    # Imprimir el formato pipeline para verificar la salida
    print(f"Formato final para el frontend (pipeline): {sentencia_list}")

    # Crear el formato esperado por el frontend basado en el pipeline
    db = get_db()
    collection = db["sentencias"]

    # Crear un pipeline para devolver los documentos que coincidan con sentencia_list
    pipeline = [
        {
            "$match": {
                "_id": {
                    "$in": [item["_id"] for item in sentencia_list if "_id" in item]
                }
            }
        },
        {"$project": {"_id": 1, "providencia": 1, "fecha_sentencia": 1}},
        {"$sort": {"fecha_sentencia": -1}},
    ]

    # Ejecutar el pipeline y devolver los resultados en el formato esperado
    sentencia_list = list(collection.aggregate(pipeline=pipeline))

    # Imprimir el formato final para verificar la salida
    print(f"Sentencia final para el frontend (pipeline): {sentencia_list}")

    return sentencia_list


def generate_evidence_checklist(text):
    match = re.search(r"```json(.*?)```", text, re.DOTALL)
    if match:
        json_text = match.group(1).strip()
    else:
        print("No JSON block found.")
        return
    # Cargar el texto JSON en un objeto
    json_object = json.loads(json_text)

    # Asegurar que cada evidencia solo tenga dos checkboxes
    for item in json_object:
        if "evidencias" in item:
            for evidencia in item["evidencias"]:
                # Añadir las opciones "Cumple" y "No cumple" por cada evidencia
                evidencia["checkboxes"] = ["Cumple", "No cumple"]

    return json_object


def get_history(user, title=""):
    db = get_db()
    collection = db["saves"]
    pipeline = []
    if title == "":
        pipeline.append({"$project": {"_id": 1, "user": 1, "modifiedAt": 1}})
        pipeline.append({"$match": {"user": user}})
        pipeline.append({"$sort": {"modifiedAt": -1}})
        pipeline.append({"$limit": 1})
    else:
        pipeline.append(
            {"$project": {"_id": 1, "user": 1, "modifiedAt": 1, "title": 1}}
        )
        pipeline.append({"$match": {"user": user, "title": title}})
    find_data = list(collection.aggregate(pipeline=pipeline))
    if len(find_data) > 1:
        return "Error"
    elif len(find_data) == 1:
        history = collection.find_one({"_id": find_data[0]["_id"]})
        print(history)
        return history
    else:
        return "No data"


def get_current_state(user):
    print(user)
    db = get_db()
    collection = db["current_state"]
    history = collection.find_one({"user": user})
    if history:
        history.pop("_id", None)
        return history
    else:
        collection.insert_one({"user": user})
        return {"user": user}


def update_current_state(user, field, data):
    db = get_db()
    collection = db["current_state"]
    query_filter = {"user": user}
    update_date = {}
    update_date[field] = data
    update_operation = {}
    update_operation["$set"] = update_date
    collection.update_one(query_filter, update_operation)


def get_current_data_field(user, field):
    db = get_db()
    collection = db["current_state"]
    history = collection.find_one({"user": user})
    if field in history:
        return history[field]
    else:
        return ""


def get_settings(field):
    db = get_db()
    collection = db["settings"]
    settings = collection.find_one()
    return settings[field]


def get_title_list(user):
    db = get_db()
    collection = db["results"]
    pipeline = [{"$project": {"title": 1, "user": 1}}, {"$match": {"user": user}}]
    res = list(collection.aggregate(pipeline=pipeline))
    return_data = []
    for each in res:
        return_data.append(each["title"])
    return return_data


def save_tutela(user, title):
    loading_data = get_current_state(user)
    loading_data["title"] = title
    loading_data["modifiedAt"] = datetime.now()

    # Asegúrate de que las evidencias se guarden correctamente
    evidencias_cumplen = session.get("evidencias_cumplen", [])
    evidencias_no_cumplen = session.get("evidencias_no_cumplen", [])

    loading_data["evidencias_cumplen"] = evidencias_cumplen
    loading_data["evidencias_no_cumplen"] = evidencias_no_cumplen

    db = get_db()
    results = db["results"]
    currents = db["current_state"]

    # Verifica si ya existe un análisis con el mismo título
    pipeline = [
        {"$project": {"_id": 1, "title": 1, "user": 1}},
        {"$match": {"user": user, "title": title}},
    ]
    res = list(results.aggregate(pipeline=pipeline))

    # Si ya existe, lo eliminamos para sobreescribir
    if len(res) > 0:
        results.delete_one({"user": user, "title": title})

    # Insertar los datos en `results`
    results.insert_one(loading_data)

    # También actualizamos el `current_state`
    currents.delete_one({"user": user})
    currents.insert_one(loading_data)

    return {"message": "sucess"}, 200


def set_tutela(user, title):
    db = get_db()
    results = db["results"]
    currents = db["current_state"]

    # Buscar los datos guardados en `results`
    set_data = results.find_one({"user": user, "title": title})

    # Si los datos se encuentran, los transferimos a `current_state`
    if set_data:
        currents.delete_one({"user": user})
        currents.insert_one(set_data)

        # Asegúrate de que las evidencias también se carguen en la sesión
        session["evidencias_cumplen"] = set_data.get("evidencias_cumplen", [])
        session["evidencias_no_cumplen"] = set_data.get("evidencias_no_cumplen", [])

        return {"message": "sucess"}, 200
    else:
        return {"message": "Análisis no encontrado"}, 404


def reset_current_state(user):
    db = get_db()
    currents = db["current_state"]
    currents.delete_one({"user": user})  # Esto debe borrar el estado anterior

    return
