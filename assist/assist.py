from pymongo import MongoClient
import os
import csv
import re
from bs4 import BeautifulSoup
import requests
import json
from flask_bcrypt import Bcrypt
from flask import Flask
import PyPDF2
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

load_dotenv()

from openai import OpenAI

OPENAI_KEY = os.getenv("OPENAI_KEY")
client2 = OpenAI(api_key=OPENAI_KEY)

client = MongoClient('mongodb://localhost:27017/tutela_db')
db = client.get_default_database()

def constdf_from_csv():
    constdf_json = []

    with open('./assist/Constdf.csv', newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            item = row['texto']
            text = item.split('.')
            articulo = text[0]
            text = articulo.split(' ')
            num = int(text[-1])
            article = articulo[:-len(text[-1])]
            texto = item[len(articulo)+2:]
            if 'transitorio' in article:
                article = 'Articulo Transitorio'
            else:
                article = 'Articulo'
            constdf_json.append({
                'num': num,
                'articulo': article,
                'texto': texto,
                'tutela': row['tutela']
            })

    print(constdf_json)

    constdf_collection = db['constdf']
    for each in constdf_json:
        constdf_collection.insert_one(each)

def sentencia_from_csv():
    sentencia_json = []

    with open('./assist/SenTencias.csv', newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            sentencia_json.append({
                'providencia': row['providencia'],
                'tipo': row['tipo'],
                'ano': row['ano'],
                'fecha_sentencia': row['fecha sentencia'],
                'tema': row['tema'],
                'magistrado': row['magistrado'],
                'fecha_publicada': row['fecha publicada'],
                'expediente': row['expediente'],
                'url': row['url'],
                'texto': "",
            })


    sentencia_collection = db['sentencias']
    for each in sentencia_json:
        sentencia_collection.insert_one(each)

def update_texto():
    sentencia_collection = db['sentencias']
    
    for x in sentencia_collection.find():
        if 'uploaded' in x and x['uploaded'] == True:
            continue
        file_path = f'../site_output/{x['providencia']}.txt'
        if os.path.isfile(file_path):
            doc = open(file_path, "r", encoding='utf-8')
            texto = doc.read()
            myquery = { "_id": x['_id'] }
            newvalues = { "$set": { "texto": texto, "uploaded": True} }
            sentencia_collection.update_one(myquery, newvalues)
            # print(f'{x['providencia']} success')
        else:
            myquery = { "_id": x['_id'] }
            newvalues = { "$set": {"uploaded": False} }
            sentencia_collection.update_one(myquery, newvalues)
            print(f'No file: {x['providencia']}')


def get_website(item):
    output_folder = '../'
    url = item['url']
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        section_div = soup.find('div', class_='Section1')
        
        if not section_div:
            section_div = soup.find('div', class_='WordSection1')
        if section_div:
            p_tag = section_div.find('p', string=lambda text: text and "Referencia:" in text)
            
            title = item['providencia'].replace('.', '-').replace('/', '-')
            title = re.sub(r'\s+', '', title)
            
            file_name = os.path.join(output_folder, f'{title}.txt')

            cleaned_content = ""
            p_tags = section_div.findAll('p')
            for p in p_tags:
                txt = p.get_text(separator=' ', strip=True)
                cleaned_text = re.sub(r'\s+', ' ', txt)
                cleaned_content = f'{cleaned_content}{cleaned_text}\n'
            with open(file_name, 'w', encoding='utf-8') as txt_file:
                txt_file.write(cleaned_content)
        else:
            print("No <div> with class 'Section1' found.", url)
            # fail_url_list.append(item)
    else:
        print("Failed to retrieve the page. Status code:", response.status_code, url)
        # fail_url_list.append(item)

    return True

def make_pipeline():
    sentencias_fields = ['providencia', 'ano', 'fecha sentencia', 'fecha publicada', 'expediente', 'url']

    keyword = "T-001"
    sortColumn = 3
    dir = 'asc'
    start = 0
    length = 10
    pipeline = []
    match_keyword = ''
    if keyword != '':
        for field in sentencias_fields:
            match_keyword = match_keyword + '{"' + field + '": {"$regex": "' + keyword + '","$options": "i"}},'
        match_keyword = match_keyword[:-1]
        or_string = '{"$or":[' + match_keyword + ']}'
        recordsFiltered = db['sentencias'].count_documents(json.loads(or_string))
        match_string = '{"$match":{"$or":[' + match_keyword + ']}}'
        pipeline.append(json.loads(match_string))

    sort_string = '{"$sort": {"' + sentencias_fields[sortColumn-1] + '":'  + str(1 if dir == 'asc' else -1) + '}}'
    pipeline.append(json.loads(sort_string))



    skip_string = '{"$skip": ' + str(start) + '}'
    pipeline.append(json.loads(skip_string))
    length_string =  '{"$limit": ' + str(length) + '}'
    pipeline.append(json.loads(length_string))

    project_string = ''
    for field in sentencias_fields:
        project_string = project_string + '"' + field + '": 1,'
    project_string = project_string[:-1]
    project_string = '{"$project":{' + project_string + '}}'
    pipeline.append(json.loads(project_string))

    # print(pipeline)

    filter_list = list(db['sentencias'].aggregate(pipeline))

    recordsTotal = db['sentencias'].count_documents({})

    print(recordsTotal)
    print(recordsFiltered)
# update_texto()

def get_sentencia():
    constitution_list = [20,30,26,35]
    collection = db['sentencias']
    pipeline = [{"$project": {"_id": 1, "providencia": 1, "fecha_sentencia": 1}},
                {"$sort": {"fecha_sentencia" : -1}},
    ]
    sentencia_list = []
    
    sentencias_list = list(collection.aggregate(pipeline=pipeline))
    for each in constitution_list:
        num = 0
        for item in sentencias_list:
            doc = collection.find_one({"_id": item['_id']})
            if fr"artículo {each}" in doc['texto']:
                sentencia_list.append(doc['providencia'])
                num += 1
            if num == 3: break
    sentencia_list = list(set(sentencia_list))
    print(sentencia_list)

def find_setencia_list(list):
    collection = db['sentencias']

    json_data = []
    for item in list:
        doc = collection.find_one({"providencia": item})
        json_data.append({
            'providencia': doc['providencia'],
            'tipo': doc['tipo'],
            'ano': doc['ano'],
            'fecha_sentencia': doc['fecha_sentencia'],
            'tema': doc['tema'],
            'magistrado': doc['magistrado'],
            'fecha_publicada': doc['fecha_publicada'],
            'expediente': doc['expediente'],
            'url': doc['url'],
        })
    print(json_data)
    return json_data

# list = ['T-243-24', 'T-245-24', 'T-232-24', 'A-1043-24', 'A-979-24', 'T-248-24', 'T-244-24', 'T-216-24']

#find_setencia_list(list)

# app = Flask(__name__)
# bcrypt = Bcrypt(app)

# def make_hash():
#     pw_hash = bcrypt.generate_password_hash('123456').decode('utf-8')
#     print(pw_hash)
# make_hash()

# file_path = 'tutela_lejandra_respuesta.pdf'
# message_file = client2.files.create(
#     file=open(file_path, "rb"), purpose="assistants"
# )

# print(message_file.id)

# client2.files.delete("file-lSyRhcfXqK5M2Mas56tOOpiN")
def file_upload_openai(file_path):
    
    message_file = client2.files.create(
        file=open(file_path, "rb"), purpose="assistants"
    )

    print(message_file.id)


def generate_evidence_checklist():
    text = '''
Visto que no se encontraron resultados de búsqueda coincidentes, enumero a continuación las evidencias necesarias basándome en el contenido del documento proporcionado:

```json
[
    {
        "descripcion": "Compra del bien inmueble ubicado en la Calle 54c sur #88I71 – Bosa (Bogotá-Colombia) el día 29 de abril de 2023",
        "evidencias": [
            {
                "tipo": "documento",
                "descripcion": "Copia de la cédula de ciudadanía del propietario del inmueble, Danilo Quevedo Vaca",
                "archivo": "cedula_danilo_quevedo.pdf"
            },
            {
                "tipo": "documento",
                "descripcion": "Documento de compra-venta del inmueble",
                "archivo": "compra_venta_inmueble.pdf"
            }
        ]
    },
    {
        "descripcion": "Elección del Sr. QUEVEDO VACA como parte del consejo de administración en el año 2017",
        "evidencias": [
            {
                "tipo": "documento",
                "descripcion": "Acta de la elección del consejo de administración",
                "archivo": "acta_eleccion_consejo_2017.pdf"
            }
        ]
    },
    {
        "descripcion": "Presentación de derecho de petición ante la administración del Conjunto Residencial Tangara Etapa 2 el 29 de abril de 2023",
        "evidencias": [
            {
                "tipo": "documento",
                "descripcion": "Copia del derecho de petición presentado ante la Administración del Conjunto Residencial Tangara Etapa 2",
                "archivo": "derecho_peticion_29042023.pdf"
            },
            {
                "tipo": "documento",
                "descripcion": "Sello de recibido confirmando la recepción del derecho de petición",
                "archivo": "sello_recibido_derecho_peticion.pdf"
            }
        ]
    },
    {
        "descripcion": "Fallos judiciales y otros actos administrativos relacionados",
        "evidencias": [
            {
                "tipo": "documento",
                "descripcion": "Sentencia de tutela a favor del solicitante",
                "archivo": "sentencia_tutela.pdf"
            },
            {
                "tipo": "documento",
                "descripcion": "Comunicaciones con la administración del conjunto residencial sobre el seguimiento de la tutela",
                "archivo": "comunicaciones_administracion.pdf"
            }
        ]
    },
    {
        "descripcion": "Visita de la Secretaría Distrital de Salud para la verificación de las instalaciones",
        "evidencias": [
            {
                "tipo": "documento",
                "descripcion": "Informe de la visita realizada por la Secretaría Distrital de Salud",
                "archivo": "informe_visita_salud.pdf"
            }
        ]
    }
]
```

Por favor, revisa el documento original o suministra los archivos mencionados para proporcionar una evaluación más precisa.
'''

    match = re.search(r'```json(.*?)```', text, re.DOTALL)
    if match:
        json_text = match.group(1).strip()
    else:
        print("No JSON block found.")
        return
    json_object = json.loads(json_text)
    
    return json_object

def make_fine_tuning_train_data():
    folder_path = '../'
    cnt = 0
    with open('output.jsonl', 'w') as f:
        for filename in os.listdir(folder_path):
            print(filename)
            if filename.endswith('.pdf'):
                if 'respuesta' in filename:
                    continue
                pdf_path1 = os.path.join(folder_path, filename)
                name, extension = filename.rsplit('.', 1)
                new_filename = f"{name}_respuesta.{extension}"
                pdf_path2 = os.path.join(folder_path, new_filename)
                try:
                    pdf_content1 = ''
                    pdf_content2 = ''
                    with open(pdf_path1, 'rb') as pdf_file:
                        reader = PyPDF2.PdfReader(pdf_file)
                        for page_num in range(len(reader.pages)):
                            page = reader.pages[page_num]
                            pdf_content1 = f'{pdf_content1}{page.extract_text()}'
                    
                    with open(pdf_path2, 'rb') as pdf_file:
                        reader = PyPDF2.PdfReader(pdf_file)
                        for page_num in range(len(reader.pages)):
                            page = reader.pages[page_num]
                            pdf_content2 = f'{pdf_content2}{page.extract_text()}'
                    cnt += 1
                    item = {"messages": [{"role": "system", "content": "Eres asistente de un abogado."}, {"role": "user", "content": f"Nuevo documento:\'{pdf_content1}\' Con base en el nuevo documento, la lista de revisiones judiciales relevantes (cuyo contenido está disponible en la tienda de vectores) y las disposiciones constitucionales proporcionadas, redacte una sentencia para el nuevo documento."}, {"role": "assistant", "content": pdf_content2}]}
                    f.write(json.dumps(item) + '\n')
                except Exception as e:
                    print(e)
        with open('output-test.jsonl', 'r', encoding='utf-8') as fi:
            for line in fi:
                json_object = json.loads(line.strip())
                f.write(json.dumps(json_object) + '\n')
                cnt += 1
                if cnt >= 10: break
            
    return

# make_fine_tuning_train_data()
def fine_tuning_upload():
    client2.files.delete("file-Og2S0kOgbIXHzVUIBpJV4AtU")
    message = client2.files.list()
    print(message.data)
    file_path = 'formato_respuesta_textanalizer.docx'
    message_file = client2.files.create(
        file=open(file_path, "rb"), purpose="assistants"
    )

    print(message_file.id)

    message = client2.fine_tuning.jobs.create(
      training_file=message_file.id,
      model="gpt-4o-mini-2024-07-18"
    )

    message = client2.fine_tuning.jobs.list()
    print(message.data)

import smtplib
def send_email_via_smtp(email, message):
    
    s = smtplib.SMTP('smtp.gmail.com', 587)
    s.starttls()
    
    MESSAGE_SENDER = str(os.getenv('MESSAGE_SENDER'))
    MESSAGE_SENDER_APP_PASSWORD = str(os.getenv('MESSAGE_SENDER_APP_PASSWORD'))
    print(email, message, MESSAGE_SENDER, MESSAGE_SENDER_APP_PASSWORD)
    s.login(MESSAGE_SENDER, MESSAGE_SENDER_APP_PASSWORD)
    # message to be sent
    # message = "Message_you_need_to_send 1"
    # sending the mail
    s.sendmail("mailer@theastralabs.com", email, message)
    # terminating the session
    s.quit()
# send_email()

def send_email_via_smtp(email, message):
    import smtplib
    import os
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    s = smtplib.SMTP('smtp.gmail.com', 587)
    s.starttls()
    MESSAGE_SENDER = str(os.getenv('MESSAGE_SENDER'))
    MESSAGE_SENDER_APP_PASSWORD = str(os.getenv('MESSAGE_SENDER_APP_PASSWORD'))
    msg = MIMEMultipart()
    msg['From'] = MESSAGE_SENDER
    msg['To'] = email
    msg['Subject'] = "Verify your email"
    
    # Attach the message body (it could be plain text or HTML)
    msg.attach(MIMEText(message, 'HTML')) 

    try:
        s.login(MESSAGE_SENDER, MESSAGE_SENDER_APP_PASSWORD)
        s.sendmail(MESSAGE_SENDER, email, msg.as_string())
        print("Email sent successfully.")
    except Exception as e:
        print(f"Failed to send email: {e}")
    finally:
        s.quit()
    return


def send_email_with_other_way(email):
    import smtplib
    import ssl
    from email.message import EmailMessage

    MESSAGE_SENDER = str(os.getenv('MESSAGE_SENDER'))
    MESSAGE_SENDER_APP_PASSWORD = str(os.getenv('MESSAGE_SENDER_APP_PASSWORD'))

    sender_email = MESSAGE_SENDER
    receiver_email = email
    password = MESSAGE_SENDER_APP_PASSWORD

    em = EmailMessage()
    em['From'] = MESSAGE_SENDER
    em['To'] = receiver_email
    em["Subject"] = "multipart test"

    # Create the plain-text and HTML version of your message
    text = """\
    Hi,
    How are you?
    Real Python has many great tutorials:
    www.realpython.com"""
    html = """\
    <html>
    <body>
        <p>Hi,<br>
        How are you?<br>
        <a href="http://www.realpython.com">Real Python</a> 
        has many great tutorials.
        </p>
    </body>
    </html>
    """

    em.set_content(html)

    context = ssl.create_default_context()
    
    print(sender_email, password)
    with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
        server.login(sender_email, password)
        server.sendmail(
            sender_email, receiver_email, em.as_string()
        )
    return

email = "crystaldev1003@gmail.com"
message = """
    <div><p>Welcome Tutela</p>
    Click <a href='http://localhost:5000/confirm_email/ImNyeXN0YWxkZXYxMDAzQGdtYWlsLmNvbSI.ZtRzgg.cCKm8J38B7L3uCTXhSH7vFLskgs'>here</a> for verify
    </div>
    """
# send_email_via_smtp(email, message)
# send_email_with_other_way(email)
# sentencia_from_csv()
# update_texto()

content = """
<p>REPÚBLICA DE COLOMBIA</p>
<p>JUZGADO [NOMBRE DEL JUZGADO]</p>
<p>[Ciudad, Fecha]</p>
<p>REFERENCIA: Acción de Tutela promovida por </p>
<h3>1. Síntesis de la sentencia</h3>
<p>La presente sentencia resuelve la Acción de Tutela presentada por Danilo Quevedo Vaca en contra de la Administración del Conjunto Residencial Tangara Etapa 2 y la Co
nstructora Compartir (Fundación Compartir). El demandante alega la vulneración de sus derechos fundamentales a la salud, a la vida, a la dignidad humana, a un ambiente
sano y a una vivienda digna, debido a problemas derivados de la ubicación inadecuada de los depósitos de basura en su lugar de residencia. El Tribunal, tras evaluar las
 pruebas y hechos presentados, ordena a las entidades demandadas tomar medidas correctivas para garantizar las condiciones de salubridad y habitabilidad adecuadas, conf
orme a las recomendaciones emitidas por la Secretaría Distrital de Salud.</p>
<h3>2. Antecedentes</h3>
<p>Danilo Quevedo Vaca, propietario de un inmueble en el Conjunto Residencial Tangara Etapa 2, ha enfrentado múltiples problemas debido a la mala ubicación y manejo de
los depósitos de basura. Estos problemas incluyen olores desagradables, la presencia de residuos líquidos y sólidos en áreas comunes, y la proliferación de bacterias, t
odo lo cual afecta gravemente la salud y la calidad de vida de los residentes. A pesar de múltiples intentos de resolver la situación a nivel administrativo y legal, la
s soluciones implementadas por la administración y la constructora han sido insuficientes.</p>
<h3>3. Trámite procesal</h3>
<p>La acción de tutela fue presentada inicialmente ante el Juez Civil de Bogotá, alegando la violación de los derechos fundamentales mencionados. La administración del
conjunto y la constructora fueron notificados y se les solicitó responder a las acusaciones. Sin embargo, la administración del conjunto no se hizo parte del proceso tu
telar en el tiempo estipulado, lo cual llevó a la emisión de una sentencia favorable al accionante, brindando protección a sus derechos.</p>
<h3>4. Hechos relevantes</h3>
<ul>
<li><strong>Primero:</strong> El 29 de abril de 2023, Danilo Quevedo Vaca compró un inmueble en el Conjunto Residencial Tangara Etapa 2. Al poco tiempo, descubrió que l
a disposición de los depósitos de basura no cumplía con estándares adecuados, afectando la salud y dignidad de los residentes.</li>
<li><strong>Segundo:</strong> Los depósitos de basura estaban ubicados a menos de 5 metros de las residencias del primer piso y permitían el libre tránsito de residuos
y olores, además de propagar bacterias.</li>
<li><strong>Tercero:</strong> En 2017, Danilo Quevedo se unió al consejo de administración del conjunto para buscar una solución al problema. La única respuesta de la c
onstructora fue instalar una puerta de madera, la cual no resolvió el problema adecuadamente.</li>
<li><strong>Cuarto:</strong> Debido a la falta de soluciones efectivas, Danilo Quevedo abandonó su residencia en 2022, lo que generó gastos adicionales ya que el inmueb
le aún no había sido totalmente pagado.</li>
<li><strong>Quinto:</strong> El 29 de abril de 2023, Danilo Quevedo presentó un derecho de petición ante la administración del conjunto, el cual no recibió respuesta ad
ecuada. Esto llevó a que presentara una acción de tutela.</li>
</ul>
<h3>5. Pruebas</h3>
<p>Entre las pruebas presentadas destacan:</p>
<ul>
<li>Copia de la cédula de ciudadanía de Danilo Quevedo Vaca.</li>
<li>Documento de queja presentado ante la administración del conjunto residencial.</li>
<li>Correspondencia por correo electrónico relacionada con la reprogramación de reuniones.</li>
<li>Videos y fotografías que documentan las condiciones inapropiadas de los depósitos de basura.</li>
<li>Histórico de correos y comunicaciones con las entidades involucradas.</li>
<li>Queja presentada por una arrendataria en marzo de 2024, evidenciando la persistencia de los problemas.</li>
</ul>
<h3>6. Decisiones objeto de revisión</h3>
<p>El Juez Civil de Bogotá declaró procedente la acción de tutela y ordenó a la Administración del Conjunto Residencial Tangara Etapa 2 y a la Constructora Compartir (F
undación Compartir) cumplir con las medidas correctivas recomendadas por la Secretaría Distrital de Salud, incluyendo el cierre seguro de los depósitos de basura y la r
ealización de limpiezas constantes en las áreas afectadas.</p>
<h3>7. Competencia</h3>
<p>La Corte Constitucional de Colombia es competente para revisar el fallo de tutela conforme a lo dispuesto en el artículo 86 de la Constitución Política y el Decreto
2591 de 1991.</p>
<h3>8. Procedencia de la acción de tutela</h3>
<h4>a) Legitimización por activa</h4>
<p>La acción de tutela fue presentada por Danilo Quevedo Vaca, quien actúa mediante su apoderado judicial, César Augusto Acosta Villabona. En este caso, el art. 86 de l
a Constitución Política de Colombia habilita a cualquier persona a interponer acción de tutela para reclamar la protección inmediata de sus derechos fundamentales cuand
o considera que estos han sido vulnerados. Danilo Quevedo Vaca, como titular de los derechos fundamentales alegadamente infringidos, está legitimado para ejercer esta a
cción.</p>
<h4>b) Legitimización por pasiva</h4>
<p>La acción de tutela se dirige contra la Administración del Conjunto Residencial Tangara Etapa 2 y la Constructora Compartir (Fundación Compartir). Estos actores son
responsables de la gestión y construcción del conjunto residencial donde ocurrieron los hechos que originaron la reclamación del tutelante. De acuerdo con la jurisprude
ncia de la Corte Constitucional, la acción de tutela procede contra particulares cuando estos actúan en posiciones de poder que puedan generar situaciones de indefensió
n o subordinación.</p>
<h4>c) Inmediatez</h4>
<p>La inmediatez es un requisito esencial en la procedencia de la acción de tutela, el cual se cumple si la solicitud se realiza dentro de un periodo prudencial desde q
ue se haga la vulneración de derechos. En este caso, Danilo Quevedo Vaca emprende la acción de tutela luego de los sucesos descritos, específicamente después de agotar
otros medios previos como la reclamación ante la Secretaría Distrital de Salud y la Administración del Conjunto Residencial sin obtener una solución efectiva. La rapide
z de la acción busca prevenir un perjuicio irremediable, en concordancia con el principio de inmediatez exigido por el art. 86 de la Constitución.</p>
<h4>d) Subsidiariedad</h4>
<p>Aunque la acción de tutela es un mecanismo subsidiario, en este caso se justifica plenamente su uso debido a la urgencia de proteger derechos fundamentales afectados
 por una acción continuada que pone en peligro la salud, la vida digna y el ambiente sano del accionante y su familia. La falta de una solución efectiva y definitiva po
r parte de la administración del conjunto residencial y de las autoridades responsables justifica la tutela como un mecanismo principal para evitar un perjuicio irremed
iable.</p>
<h3>9. Delimitación del caso, formulación del problema jurídico y esquema de la decisión</h3>
<h4>a) Delimitación del caso</h4>
<p>El caso se centra en la presunta vulneración de los derechos fundamentales de Danilo Quevedo Vaca y su núcleo familiar por la indebida ubicación y mantenimiento de l
os depósitos de basura del Conjunto Residencial Tangara Etapa 2, lo cual ha ocasionado graves problemas de salubridad y dignidad humana, afectando también a terceras pe
rsonas, como la arrendataria del inmueble, su madre de tercera edad y su hijo menor.</p>
<h4>b) Formulación del problema jurídico</h4>
<p>El problema jurídico se enfoca en determinar si se han vulnerado los derechos fundamentales de Danilo Quevedo Vaca a la salud, vida digna y vivienda digna por la con
ducta omisiva de la Administración del Conjunto Residencial Tangara Etapa 2 y Constructora Compartir, quienes no han corregido adecuadamente los problemas estructurales
 relacionados con los residuos sólidos pese a las múltiples reclamaciones y recomendaciones de la Secretaría Distrital de Salud.</p>
<h4>c) Esquema de la decisión</h4>
<p>Para resolver este problema, se deberá analizar la responsabilidad de la administración y constructora en la calidad de vida de los residentes, reiterar la jurisprud
encia relacionada con la garantía de derechos fundamentales en el marco de la propiedad horizontal y evaluar las evidencias aportadas que demuestran la persistencia de
los problemas pese a las medidas superficiales tomadas.</p>
<h3>10. Solución del caso concreto</h3>
<p>Al analizar el caso concreto, se verifica que la administración del conjunto residencial y la constructora han incurrido en omisiones que han resultado en la vulnera
ción de los derechos fundamentales del accionante y otras personas afectadas. La respuesta insuficiente a las reclamaciones, la falta de implementación de soluciones du
raderas y la continua exposición a condiciones insalubres constituyen una clara violación de los principios constitucionales de dignidad humana y ambiente sano.</p>
<p>Se instruye a las partes demandadas a:</p>
<ol>
<li>Implementar de inmediato soluciones técnicas definitivas que eliminen el problema de disposición de residuos de acuerdo con las recomendaciones emitidas por la Secr
etaría Distrital de Salud.</li>
<li>Realizar inspecciones y mantener de manera constante la higiene de las áreas comunes afectadas.</li>
<li>Ofrecer compensaciones adecuadas a los residentes afectados por las condiciones insalubres vividas durante el periodo de la omisión.</li>
</ol>
<h3>11. Problema jurídico</h3>
<p>El problema jurídico del presente caso se fundamenta en establecer si las condiciones de insalubridad generadas por la incorrecta ubicación y mantenimiento de los de
pósitos de basura del Conjunto Residencial Tangara Etapa 2 vulneran el derecho a la vida, a la salud, y a la vivienda digna del señor Danilo Quevedo Vaca y su núcleo fa
miliar. Y si, ante la omisión de la administración y constructora en solucionar de manera efectiva y definitiva estos problemas, procede la acción de tutela como mecani
smo idóneo para proteger estos derechos fundamentales.</p>
<h3>12. Consideraciones</h3>
<p>En el presente caso, se observa que la problemática central gira en torno a la ubicación y gestión inadecuada de los depósitos de basura en el Conjunto Residencial T
angara Etapa 2, que ha afectado de manera significativa la salud y dignidad de los residentes, especialmente del señor Danilo Quevedo Vaca y su familia. La Acción de Tu
tela interpuesta busca la protección de los derechos fundamentales vulnerados debido a esta situación.</p>
<h4>Derecho a la Salud y un Ambiente Sano</h4>
<p>El derecho a la salud está consagrado en la Constitución Nacional y en diversas jurisprudencias de la Corte Constitucional. Dicho derecho implica no solo la ausencia
 de enfermedades, sino también la existencia de condiciones ambientales adecuadas para una vida digna. En este caso, la indebida colocación de los depósitos de basura h
a propiciado la propagación de residuos y bacterias, afectando las áreas comunes y poniendo en riesgo la salud de los residentes.</p>
<h4>Derecho a la Dignidad Humana</h4>
<p>La Corte Constitucional ha reiterado en múltiples fallos, como la Sentencia T-444/99, que el derecho a la vida implica garantizar una existencia digna, la cual abarc
a tanto el desarrollo físico como espiritual del individuo. En este sentido, cualquier acción u omisión que perjudique seriamente la calidad de vida de una persona podr
ía considerarse una violación a este derecho primordial.</p>
<h4>Vulneración de Derechos Fundamentales</h4>
<p>La situación descrita por el accionante evidencia una vulneración clara de derechos fundamentales como la vida, la salud, un ambiente sano y la dignidad humana. La f
alta de una respuesta adecuada por parte de la Administración del Conjunto Residencial Tangara Etapa 2 y de la Constructora Compartir indica una omisión irresponsable,
que resulta en un daño continuo a las condiciones de vida de los residentes.</p>
<h3>13. Decisión</h3>
<p>En virtud de las consideraciones precedentemente expuestas, este despacho judicial considera procedente amparar los derechos invocados por el señor Danilo Quevedo Vaca. Se concluye que ha existido una evidente vulneración a sus derechos fundamentales debido a la ubicación inadecuada y la gestión deficiente de los residuos en su lug
ar de residencia.</p>
<h3>14. Resuelve</h3>
<p>Primero: <strong>Tutelar</strong> los derechos fundamentales a la salud, vida, dignidad humana, ambiente sano y vivienda digna del señor Danilo Quevedo Vaca, su fami
lia y demás residentes del Conjunto Residencial Tangara Etapa 2.</p>
<p>Segundo: <strong>Ordenar</strong> a la Constructora Compartir (Fundación Compartir) que, dentro de los términos de garantía, acate y ejecute las recomendaciones impa
rtidas por la Secretaría Distrital de Salud respecto a la reubicación y correcta gestión de los depósitos de basura.</p>
<p>Tercero: <strong>Ordenar</strong> a la Administración del Conjunto Residencial Tangara Etapa 2 y su Consejo de Administración que cumplan de manera inmediata con las
 recomendaciones de la Secretaría Distrital de Salud y adopten las medidas correctivas necesarias para solucionar de fondo la problemática presentada.</p>
<p>Cuarto: <strong>Solicitar</strong> a la Defensoría del Pueblo que supervise y acompañe el cumplimiento de esta sentencia, garantizando la efectiva protección de los
derechos de los involucrados.</p>
<p>Quinto: <strong>Comunicar</strong> esta decisión a las partes involucradas y a las autoridades competentes para su debido cumplimiento.</p>
<p>Cópiese, notifíquese y cúmplase.</p>
<p>[Nombre del Juez]Notas: </p>
<p>Revisión de la jurisprudencia relevante: </p>
<p>T-416-01</p><p>T-881-02</p>
<p>T-924A-13</p>"""


def add_html_to_docx(soup, doc):
    for element in soup:
        if '\n' == element:
            continue
        print(element)
        if element.name == 'p':
            paragraph = doc.add_paragraph(element.get_text())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)   # 0.5 line space after
        elif element.name == 'h1':
            paragraph = doc.add_heading(level=1)
            run = paragraph.add_run(element.get_text())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)   # 0.5 line space after
        elif element.name == 'h2':
            paragraph = doc.add_heading(level=2)
            run = paragraph.add_run(element.get_text())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)   # 0.5 line space after
        elif element.name == 'h3':
            paragraph = doc.add_heading(level=3)
            run = paragraph.add_run(element.get_text())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)   # 0.5 line space after
        elif element.name == 'strong':
            run = doc.add_paragraph().add_run(element.get_text())
            paragraph_format = run.paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)   # 0.5 line space after
        elif element.name == 'em':
            run = doc.add_paragraph().add_run(element.get_text())
            run.italic = True
            paragraph_format = run.paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)   # 0.5 line space after
        elif element.name == 'ul':
            list_items = element.find_all('li')
            for i, li in enumerate(list_items):
                paragraph = doc.add_paragraph(li.get_text(), style='ListBullet')
                paragraph_format = paragraph.paragraph_format
                if i == 0:
                    paragraph_format.space_before = Pt(6)
                if i == len(list_items) - 1:
                    paragraph_format.space_after = Pt(6)
        elif element.name == 'ol':
            list_items = element.find_all('li')
            for i, li in enumerate(list_items):
                paragraph = doc.add_paragraph(li.get_text(), style='ListNumber')
                paragraph_format = paragraph.paragraph_format
                if i == 0:
                    paragraph_format.space_before = Pt(6)
                if i == len(list_items) - 1:
                    paragraph_format.space_after = Pt(6)
                    
        elif isinstance(element, str):
            paragraph = doc.add_paragraph(element)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)  # 0.5 line space before
            paragraph_format.space_after = Pt(6)   # 0.5 line space after

def create_docx_from_html(html_content, filename):
    soup = BeautifulSoup(html_content, 'lxml')

    doc = Document()
    add_html_to_docx(soup.body.contents, doc)
    
    doc.save(filename)

create_docx_from_html(content, 'a.docx')